import google.generativeai as genai
import os
import time
import pandas as pd
from google.generativeai import caching
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import json
import re
import xml.etree.ElementTree as ET

# Load environment variables from .env file
load_dotenv()

# Configure API Key (Ensure to set this in your environment)
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

def verify_video(video_path):
    if not os.path.exists(video_path):
        raise FileNotFoundError("Video file not found.")
    if not video_path.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
        raise ValueError("Invalid file format. Only MP4, AVI, MOV, and MKV videos are allowed.")
    return True

def upload_to_gemini(path, mime_type="video/mp4"):
    file = genai.upload_file(path, mime_type=mime_type)
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file

def wait_for_files_active(files):
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
    print("Processing complete.")

def set_table_borders(table, border_color="auto"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = OxmlElement(f'w:{border}')
                border_elem.set(qn('w:val'), 'single')
                border_elem.set(qn('w:sz'), '4')
                border_elem.set(qn('w:space'), '0')
                if border_color != "auto":
                    border_elem.set(qn('w:color'), border_color)
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)

def process_video(video_path):
    verify_video(video_path)
    file = upload_to_gemini(video_path)
    wait_for_files_active([file])
    model = genai.GenerativeModel("gemini-1.5-pro-latest")
    response = model.generate_content(["""You are a Business Analyst tasked with reviewing a process recording from the Subject Matter Expert (SME) in the form of a video. Your objective is to carefully analyze the video and extract a detailed, step-by-step outline of the process presented. The video may not cover the process end-to-end, so you need to assess both the explicit steps presented and any references the SME makes to previous steps.

Your outline should be clear, precise, and suitable for inclusion in formal documentation, such as a Process Definition Document (PDD). Ensure that each step is detailed, any business exceptions are noted, and the process is presented in the order it is executed. Pay attention to the narrator’s comments to identify any transitions or additional information.

The structure of the output documentation should include the following sections:

1. Process Name
   Provide the name of the process being described.

2. Short Process Description
   Offer a brief summary of the process.

3. List of Applications Utilized
   This should be a table that includes the following details for each application used in the process:
   - The name of the application
   - The type of the application (e.g., web application, desktop application)
   - The URL of the application, if applicable
   Ensure both web and desktop applications are identified.

4. List of Steps
   - Provide a detailed, step-by-step description of the process in the order the steps are executed.
   - Steps should be listed as they were presented in the video.
   - Each interaction with the user interface (UI) must be documented.
   - Document each described or presented data transformation.
   - Use the following numbering format:
     - Example:
       1.0 Group of steps
       1.1 First step in the group
       1.2 Second step in the group
   - Steps should specify the UI element the user interacts with or the calculation logic described.
   - First step the group should specify the application name that the user interact with.

5. Exception Handling
   Describe any exceptions in the process and how they should be handled.

6. Requires Clarification
   List any questions you have for the SME or aspects of the process that require further clarification.

Provide the output in the following JSON format:

{
  "process_name": "[The name of the process based on the video content]",
  "short_process_description": "[The short process description based on the video content]",
  "list_of_applications": [
    {
      "application_name": "[Name of the application]",
      "type": "[Type of the application, e.g., web/desktop]",
      "url": "[URL of the application, if applicable]"
    },
    {
      "application_name": "[Name of the application]",
      "type": "[Type of the application, e.g., web/desktop]",
      "url": "[URL of the application, if applicable]"
    }
  ],
  "list_of_steps": [
    {
      "group_name": "[Description of the group of steps]",
      "numbering": "1.0",
      "time_stamp": "[Timestamp from the video when this step is executed]",
      "sub_steps": [
        {
          "step": "[Description of the sub-step]",
          "numbering": "1.1",
          "time_stamp": "[Timestamp from the video when this step is executed]"
        },
        {
          "step": "[Description of the sub-step]",
          "numbering": "1.2",
          "time_stamp": "[Timestamp from the video when this step is executed]"
        }
      ]
    }
  ],
  "exceptions": [
    {
        "exception": "[Exception name]",
        "description": "[Exception description]"
    },
    {
        "exception": "[Exception name]",
        "description": "[Exception description]"
    },
    {
        "exception": "[Exception name]",
        "description": "[Exception description]"
    }
  ],
  "clarifications": [
    "[Required clarification or question]",
    "[Required clarification or question]",
    "[Required clarification or question]"
  ]
}""", file])
    
    processed_text = response.text if response else "No data extracted."
    pattern = r"\{.*\}"
    json_match = re.search(pattern, processed_text, re.DOTALL)
    if json_match:
        json_str = json_match.group(0)
        try:
            json_data = json.loads(json_str)
        except json.JSONDecodeError as e:
            print("Error decoding JSON:", e)
            return None, None
    else:
        print("No JSON found in the text.")
        return None, None

    # Generate Word file
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading('Process Name: ' + json_data["process_name"], level=1)
    doc.add_paragraph(json_data["short_process_description"])
    doc.add_heading('List of applications', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Application Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'URL'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    for app in json_data["list_of_applications"]:
        row_cells = table.add_row().cells
        row_cells[0].text = app['application_name']
        row_cells[1].text = app['type']
        row_cells[2].text = app['url'] if app['url'] is not None else ''
    set_table_borders(table)
    doc.add_heading('List of steps', level=2)
    for step_group in json_data["list_of_steps"]:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        col_widths = [Inches(0.5), Inches(4.5), Inches(1.5)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width
        hdr_cells = table.rows[0].cells
        hdr_cell = hdr_cells[0].merge(hdr_cells[2])
        hdr_cell.paragraphs[0].add_run(f"{step_group['numbering']} {step_group['group_name']}").bold = True
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D1E2F8')
        hdr_cell._element.tcPr.append(shading_elm)
        for sub_step in step_group["sub_steps"]:
            row_cells = table.add_row().cells
            row_cells[0].text = sub_step['numbering']
            row_cells[1].text = sub_step['step']
            row_cells[2].text = sub_step['time_stamp']
        set_table_borders(table)
    word_file_path = f"{json_data['process_name'].replace(' ', '_')}.docx"
    doc.save(word_file_path)
    print("Word file generated:", word_file_path)

   # Generate Draw.io file
    systemPrompt_ChatGeneration="""Generate the MX file code that can be directly imported into Draw.io to create a process map. The MX should define a workflow process map according the <List_of_Steps/> in the user's message, including shapes (e.g., rectangles for actions, diamonds for decisions), connectors, and labels. Make sure that the MX file generated by you is correct and can be copied to the Draw.IO application without errors.

Follow the rules below:
<RULES>
Start and End nodes: Represent the beginning and end of the process.
Reflect on the chart all the steps listed in the 'List_of_Steps' section.
Do not mearge steps from the list of steps to the one block (mxCell) in the MX file.
Do not group steps under each other.
The order of steps on the chart should be from left to right.
There are no space limitations. Order blocks horizontally.
Action nodes: Use rectangles to represent tasks or actions in the process.
Decision nodes: Use diamonds to represent decision points with 'Yes' and 'No' branches.
Group of steps are logically connected thus all steps needs to be in one line and not grouped.
<IMPORTANT> The text inside value tag must not contain double quotes "" use singel quotes instead. '' ex: value=" Product name is: 'Name of the Product'" or value="Click 'Submit' button"! <IMPORTANT/>
Each action listed on the 'List of Steps' needs to be a separate step (block)
Connectors: Connect shapes to define the flow from one step to the next.
Labels: Add labels to each shape to describe its role (e.g., 'Start', 'Decision: Is data valid?', 'End').
Do not use double quates in the value for the MX elements.
Do not use double quates inside the value tag. Use Single quates instead.
The flow of the chart should be fromleft to right.
The diagram should not have any bended lines.
Generate the MX file till all the steps are reflected and the file has closuing clause of "</diagram></mxfile>"
Start ad End nodes should be a circle (not elipses).
The list of steps should be as detailed as possible.
Make sure taht steps are propery connected.
Make sure that all the items are properly alligend.
Remember to add any rules or decision points using dimonds.
Remember to split each point into the separate box.
</RULES>

<IMPORTAINT>
Be detailed and do not miss any steps in the <List_of_Steps/>.  The content will be copied and pasted to the blank Draw.IO project. Make sure that the XLM is correcly structed and will work. Please find the exmaple of the correct Draw IO format below. Do not define 'etag'.

<CRITICAL> The text inside value tag must not contain double quotes "" use singel quotes instead. '' ex: value=" Product name is: 'Name of the Product'" or value="Click 'Submit' button"! <IMPORTANT/>
</CRITICAL>

<MX FILE EXAMPLE>

<mxfile host="Electron" modified="2024-11-11T23:25:00.048Z" agent="5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) draw.io/20.2.3 Chrome/102.0.5005.167 Electron/19.0.11 Safari/537.36" etag="SEY0rUoS_tAO2XkjiwVg" compressed="false" version="20.2.3" type="device">
  <diagram id="Yt7LL3NTDxsKtHlv3eUc" name="[Process Name]">
    <mxGraphModel dx="3000" dy="3000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="3000" pageHeight="3000" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        <mxCell id="Cmv73IDUztHTh19W1OSw-33" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" target="Cmv73IDUztHTh19W1OSw-36" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-65" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0;entryY=0.5;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-34" target="Cmv73IDUztHTh19W1OSw-36" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-34" value="Start" style="ellipse;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="340" y="860" width="60" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-35" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-36" target="Cmv73IDUztHTh19W1OSw-38" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-36" value="Detailed step one" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="440" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-37" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-38" target="Cmv73IDUztHTh19W1OSw-40" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-38" value="Detailed step two" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="600" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-39" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-40" target="Cmv73IDUztHTh19W1OSw-44" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-40" value="Detailed step &#39;three&#39;" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="760" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-41" value="Yes" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-44" target="Cmv73IDUztHTh19W1OSw-46" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-42" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0;entryY=0.5;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-44" target="Cmv73IDUztHTh19W1OSw-48" edge="1">
          <mxGeometry relative="1" as="geometry">
            <Array as="points">
              <mxPoint x="960" y="980" />
            </Array>
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-43" value="No" style="edgeLabel;html=1;align=center;verticalAlign=middle;resizable=0;points=[];" parent="Cmv73IDUztHTh19W1OSw-42" vertex="1" connectable="0">
          <mxGeometry x="0.1738" relative="1" as="geometry">
            <mxPoint as="offset" />
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-44" value="Condition" style="rhombus;whiteSpace=wrap;html=1;rounded=0;" parent="1" vertex="1">
          <mxGeometry x="920" y="850" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-45" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-46" target="Cmv73IDUztHTh19W1OSw-50" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-46" value="Decision A" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1070" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-47" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-48" target="Cmv73IDUztHTh19W1OSw-53" edge="1">
          <mxGeometry relative="1" as="geometry">
            <Array as="points">
              <mxPoint x="1130" y="1150" />
              <mxPoint x="1930" y="1150" />
            </Array>
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-48" value="Decision B" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1070" y="950" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-49" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-50" target="Cmv73IDUztHTh19W1OSw-58" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-50" value="Detailed step &#39;three&#39;" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1230" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-51" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-52" target="Cmv73IDUztHTh19W1OSw-53" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-52" value="Detailed step" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1740" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-53" value="End" style="ellipse;whiteSpace=wrap;html=1;rounded=0;" parent="1" vertex="1">
          <mxGeometry x="1900" y="860" width="60" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-54" value="Outcome A" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;exitX=1;exitY=0.5;exitDx=0;exitDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-58" target="Cmv73IDUztHTh19W1OSw-60" edge="1">
          <mxGeometry relative="1" as="geometry">
            <mxPoint x="1500.0000000000005" y="890" as="sourcePoint" />
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-55" value="Outcome B" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0;entryY=0.5;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-58" target="Cmv73IDUztHTh19W1OSw-62" edge="1">
          <mxGeometry x="0.2646" relative="1" as="geometry">
            <Array as="points">
              <mxPoint x="1440" y="990" />
            </Array>
            <mxPoint as="offset" />
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-56" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0;entryY=0.5;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-58" target="Cmv73IDUztHTh19W1OSw-64" edge="1">
          <mxGeometry relative="1" as="geometry">
            <Array as="points">
              <mxPoint x="1440" y="1090" />
            </Array>
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-57" value="Outcome C" style="edgeLabel;html=1;align=center;verticalAlign=middle;resizable=0;points=[];" parent="Cmv73IDUztHTh19W1OSw-56" vertex="1" connectable="0">
          <mxGeometry x="0.4742" y="1" relative="1" as="geometry">
            <mxPoint as="offset" />
          </mxGeometry>
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-58" value="Decision" style="rhombus;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1400" y="850" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-59" value="" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;" parent="1" source="Cmv73IDUztHTh19W1OSw-60" target="Cmv73IDUztHTh19W1OSw-52" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-60" value="Detailed step" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1580" y="860" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-61" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0.5;entryY=1;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-62" target="Cmv73IDUztHTh19W1OSw-53" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-62" value="Detailed step" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1580" y="960" width="120" height="60" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-63" style="edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;entryX=0.5;entryY=1;entryDx=0;entryDy=0;" parent="1" source="Cmv73IDUztHTh19W1OSw-64" target="Cmv73IDUztHTh19W1OSw-53" edge="1">
          <mxGeometry relative="1" as="geometry" />
        </mxCell>
        <mxCell id="Cmv73IDUztHTh19W1OSw-64" value="Detailed step" style="rounded=0;whiteSpace=wrap;html=1;" parent="1" vertex="1">
          <mxGeometry x="1580" y="1060" width="120" height="60" as="geometry" />
        </mxCell>
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>"""
    generation_config = {
        "temperature": 1,
        "top_p": 0.8,
        "top_k": 40,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
    list_of_steps = json.dumps(json_data["list_of_steps"], indent=4)
    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash-002",
        generation_config=generation_config,
        system_instruction=systemPrompt_ChatGeneration
    )
    chat_session = model.start_chat()
    response = chat_session.send_message(list_of_steps)
    ChartGeneratedByAI = response.text

    # Generate file name based on the process_name
    file_name = f"{json_data['process_name'].replace(' ', '_')}.drawio"
    pattern = r"```xml(.*?)```"
    match = re.search(pattern, ChartGeneratedByAI, re.DOTALL)
    if match:
        xml_content = match.group(1).strip()
    else:
        print("No XML content found.")
        return word_file_path, None
    file_path = f"./{file_name}"
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(xml_content)
    print(f"Draw.io file saved as {file_name}")
    
    return word_file_path, file_path
